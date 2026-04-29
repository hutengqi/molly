"""Generate a poi-tl compatible contract.docx template.

Placeholders follow poi-tl syntax: {{var}}, {{?cond}}{{/cond}}, {{*list}}{{/list}}.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "molly-document-example/src/main/resources/templates/word/contract.docx"

doc = Document()

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# 全局字体（中文建议 PingFang / 宋体；这里交给系统解释）
style = doc.styles["Normal"]
style.font.name = "SimSun"
style.font.size = Pt(11)

# 标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("服 务 合 同")
run.bold = True
run.font.size = Pt(20)

# 合同编号
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
meta.add_run("合同编号：{{contractNo}}")

# 签订信息
p = doc.add_paragraph()
p.add_run("甲方（委托方）：").bold = True
p.add_run("{{partyA}}")
p = doc.add_paragraph()
p.add_run("乙方（服务方）：").bold = True
p.add_run("{{partyB}}")
p = doc.add_paragraph()
p.add_run("签订日期：").bold = True
p.add_run("{{signDate}}")

doc.add_paragraph()

# 条款
doc.add_heading("一、服务内容", level=2)
doc.add_paragraph("乙方根据甲方委托，向甲方提供 {{serviceName}} 服务，服务说明：{{serviceDesc}}。")

doc.add_heading("二、服务金额", level=2)
doc.add_paragraph("本合同服务总金额为：人民币 {{amount}} 元（大写：{{amountCn}}）。")

doc.add_heading("三、服务明细", level=2)
table = doc.add_table(rows=4, cols=4)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "序号"
hdr[1].text = "项目名称"
hdr[2].text = "数量"
hdr[3].text = "金额"
for i in range(1, 4):
    cells = table.rows[i].cells
    cells[0].text = str(i)
    cells[1].text = f"{{{{item{i}Name}}}}"
    cells[2].text = f"{{{{item{i}Qty}}}}"
    cells[3].text = f"{{{{item{i}Price}}}}"

doc.add_heading("四、付款方式", level=2)
doc.add_paragraph("{{paymentTerms}}")

doc.add_heading("五、违约责任", level=2)
doc.add_paragraph("任何一方违反本合同约定，应向守约方支付违约金，金额为未履行部分合同金额的 20%。")

doc.add_heading("六、签署", level=2)
signt = doc.add_table(rows=2, cols=2)
signt.rows[0].cells[0].text = "甲方（盖章）："
signt.rows[0].cells[1].text = "乙方（盖章）："
signt.rows[1].cells[0].text = "授权代表：{{partyASigner}}"
signt.rows[1].cells[1].text = "授权代表：{{partyBSigner}}"

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
foot.add_run("—— 本合同由 molly-document-spring-boot-starter 渲染生成 ——").italic = True

doc.save(OUT)
print(f"wrote {OUT}")
